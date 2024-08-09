import io
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import pdfplumber
import pandas as pd
import docx
from odf import text, teletype
from odf.opendocument import load


def process_file(file_content, file_extension):
    if file_extension == 'pdf':
        return pdf_service(file_content)
    elif file_extension in ['csv', 'xlsx', 'xls']:
        return spreadsheet_service(file_content, file_extension)
    elif file_extension in ['doc', 'docx']:
        return word_service(file_content)
    elif file_extension == 'odt':
        return odt_service(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def pdf_service(pdf_file):
    text_content = ""
    tables = []

    # PyMuPDF ile PDF açma
    doc = fitz.open(stream=pdf_file, filetype="pdf")

    for page_num in range(len(doc)):
        page = doc[page_num]

        # Metin çıkarma
        text_content += page.get_text()

        # Resimlerden metin çıkarma (OCR)
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            # PIL kullanarak resmi açma
            image = Image.open(io.BytesIO(image_bytes))

            # Pytesseract ile OCR işlemi
            text_content += pytesseract.image_to_string(image)

    # pdfplumber ile tablo çıkarma
    with pdfplumber.open(io.BytesIO(pdf_file)) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            tables.extend(page_tables)

    return {"text": text_content, "tables": tables}


def spreadsheet_service(file_content, file_extension):
    if file_extension == 'csv':
        df = pd.read_csv(io.BytesIO(file_content))
    else:  # xlsx or xls
        df = pd.read_excel(io.BytesIO(file_content))

    text_content = df.to_string(index=False)
    tables = [df.values.tolist()]

    return {"text": text_content, "tables": tables}


def word_service(file_content):
    doc = docx.Document(io.BytesIO(file_content))
    text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    tables = [[cell.text for cell in row.cells] for table in doc.tables for row in table.rows]

    return {"text": text_content, "tables": tables}


def odt_service(file_content):
    doc = load(io.BytesIO(file_content))
    text_content = ""
    for element in doc.getElementsByType(text.P):
        text_content += teletype.extractText(element) + "\n"

    tables = []
    for table in doc.getElementsByType(text.Table):
        table_data = []
        for row in table.getElementsByType(text.TableRow):
            row_data = []
            for cell in row.getElementsByType(text.TableCell):
                cell_content = ""
                for p in cell.getElementsByType(text.P):
                    cell_content += teletype.extractText(p) + " "
                row_data.append(cell_content.strip())
            table_data.append(row_data)
        tables.append(table_data)

    return {"text": text_content, "tables": tables}